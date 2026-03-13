#!/usr/bin/env python3
"""
PDF论文分析工具 - 主脚本
功能：自动提取PDF论文中的关键信息并生成Excel报告
支持 PaddleOCR 文本识别，适用于扫描版PDF
作者：WorkBuddy AI Assistant
版本：v1.4
"""

from __future__ import annotations

import argparse
import csv
import json
import re
from datetime import datetime
from pathlib import Path
from typing import Optional

import openpyxl
import pdfplumber
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

from ocr_engine import create_ocr_engine, OCREngine

# ============================================================================
# 配置区域
# ============================================================================

# 项目根目录
PROJECT_ROOT = Path(__file__).resolve().parent.parent

# PDF文件所在目录（用户可以修改）
PDF_DIR = PROJECT_ROOT / "示例数据" / "pdfs"

# 输出目录（用户可以修改）
OUTPUT_DIR = PROJECT_ROOT / "输出结果"

# 是否使用已有的JSON数据（如果存在）
USE_EXISTING_JSON = True

# 期刊影响因子JSON配置文件
JOURNAL_CONFIG_PATH = PROJECT_ROOT / "配置文件" / "journal_impact_factors.json"

# Excel/导出字段顺序
REPORT_HEADERS = [
    "File",
    "URL",
    "期刊名称",
    "影响因子",
    "作者",
    "论文标题",
    "器件结构",
    "优化层级",
    "优化策略",
    "EQE（外量子效率）",
    "色度坐标",
    "寿命",
    "补充信息",
]

# ============================================================================
# 期刊影响因子数据库（2023-2024年数据，用户可自定义添加）
# ============================================================================

JOURNAL_IMPACT_FACTORS = {
    # 顶级期刊
    'Nature Communications': 16.6,
    'Nature': 64.8,
    'Nature Photonics': 39.7,
    'Nature Materials': 41.2,
    'Nature Nanotechnology': 38.3,
    'Science': 56.9,
    
    # 高影响因子期刊
    'Advanced Materials': 29.4,
    'Advanced Functional Materials': 19.0,
    'Advanced Optical Materials': 10.0,
    'ACS Nano': 17.1,
    'Nano Letters': 12.8,
    'Nano Today': 18.8,
    'Small': 13.3,
    'Small Methods': 12.4,
    
    # 化学和材料期刊
    'Journal of the American Chemical Society': 15.0,
    'JACS': 15.0,
    'ACS Applied Materials & Interfaces': 9.5,
    'Chemical Engineering Journal': 15.1,
    'ACS Applied Electronic Materials': 5.3,
    
    # 物理和光学期刊
    'Laser & Photonics Reviews': 10.9,
    'Optics Express': 3.8,
    'Optics Letters': 3.6,
    'Journal of Physical Chemistry C': 3.7,
    
    # 其他期刊（用户可添加）
    'Materials Today': 24.2,
    'Nano Research': 10.3,
    'Journal of Materials Chemistry C': 6.4,
    'Dalton Transactions': 4.1,
    'Physical Chemistry Chemical Physics': 3.7,
    'Journal of Photochemistry and Photobiology C': 12.1,
}

# ============================================================================
def load_journal_impact_factors(config_path, use_existing_json=True):
    """加载期刊影响因子配置"""
    impact_factors = dict(JOURNAL_IMPACT_FACTORS)

    if not use_existing_json or not config_path.exists():
        return impact_factors

    try:
        with config_path.open("r", encoding="utf-8") as file:
            config = json.load(file)
    except (OSError, json.JSONDecodeError) as exc:
        print(f"[警告] 无法读取期刊配置文件 {config_path}: {exc}")
        return impact_factors

    for category, journals in config.items():
        if category == "说明" or not isinstance(journals, dict):
            continue

        for journal_name, info in journals.items():
            if isinstance(info, dict) and "impact_factor" in info:
                impact_factors[journal_name] = info["impact_factor"]
            elif isinstance(info, (int, float)):
                impact_factors[journal_name] = float(info)

    return impact_factors


def parse_args():
    """解析命令行参数"""
    parser = argparse.ArgumentParser(
        description="批量分析 PDF 论文并导出 Excel 报告。"
    )
    parser.add_argument(
        "--pdf-dir",
        type=Path,
        default=PDF_DIR,
        help=f"PDF 输入目录，默认: {PDF_DIR}",
    )
    parser.add_argument(
        "--recursive",
        action="store_true",
        help="递归扫描子目录内的 PDF 文件。",
    )
    parser.add_argument(
        "--output-dir",
        type=Path,
        default=OUTPUT_DIR,
        help=f"Excel 输出目录，默认: {OUTPUT_DIR}",
    )
    parser.add_argument(
        "--journal-config",
        type=Path,
        default=JOURNAL_CONFIG_PATH,
        help=f"期刊影响因子 JSON 配置文件，默认: {JOURNAL_CONFIG_PATH}",
    )
    parser.add_argument(
        "--no-json",
        action="store_true",
        help="不加载 JSON 期刊配置，仅使用脚本内置期刊数据。",
    )
    parser.add_argument(
        "--max-pages",
        type=int,
        default=0,
        help="每篇论文最多读取的页数（0 表示不限制）。用于加速但可能降低提取命中率。",
    )
    parser.add_argument(
        "--export-json",
        action="store_true",
        help="除 Excel 外，同时导出 JSON 结果到输出目录。",
    )
    parser.add_argument(
        "--export-csv",
        action="store_true",
        help="除 Excel 外，同时导出 CSV 结果到输出目录。",
    )
    parser.add_argument(
        "--export-md",
        action="store_true",
        help="除 Excel 外，同时导出 Markdown 缺失项报告到输出目录。",
    )
    parser.add_argument(
        "--export-docx",
        "--export-word",
        dest="export_docx",
        action="store_true",
        help="除 Excel 外，同时导出 Word(.docx) 缺失项报告到输出目录。",
    )
    parser.add_argument(
        "--use-ocr",
        action="store_true",
        default=False,
        help="使用 PaddleOCR 进行文本提取（强制模式）。",
    )
    parser.add_argument(
        "--no-ocr",
        action="store_true",
        help="禁用 OCR，使用 pdfplumber 提取文本（适用于有文本层的PDF）。",
    )
    parser.add_argument(
        "--ocr-lang",
        type=str,
        default="en",
        choices=["en", "ch", "japan", "korean", "fr", "de", "ru"],
        help="OCR 识别语言，默认: en",
    )
    parser.add_argument(
        "--enable-hpi",
        action="store_true",
        help="启用 PaddleOCR 高性能推理模式（需要额外安装依赖）。",
    )
    return parser.parse_args()


# 核心功能函数
# ============================================================================

def iter_pdf_files(pdf_path_or_dir: Path, recursive: bool = False) -> list[Path]:
    """收集待分析的 PDF 文件列表（支持目录/单文件，支持 .pdf / .PDF 等大小写）"""
    path = pdf_path_or_dir.expanduser()
    if not path.exists():
        return []

    if path.is_file():
        return [path] if path.suffix.lower() == ".pdf" else []

    if recursive:
        candidates = (p for p in path.rglob("*") if p.is_file())
    else:
        candidates = (p for p in path.iterdir() if p.is_file())

    pdfs = [p for p in candidates if p.suffix.lower() == ".pdf"]
    pdfs.sort(key=lambda p: str(p).lower())
    return pdfs


def extract_pdf_text(
    pdf_path: Path,
    max_pages: Optional[int] = None,
    ocr_engine: Optional[OCREngine] = None,
):
    """提取PDF文本与元数据

    Args:
        pdf_path: PDF文件路径
        max_pages: 最大读取页数
        ocr_engine: OCR引擎实例（如果提供则使用OCR模式）

    返回: (full_text, front_text, metadata)
      - full_text: 多页拼接后的全文（可能受 max_pages 限制）
      - front_text: 第一段可用的"前置信息页"文本（用于期刊/标题/作者识别）
      - metadata: PDF 元数据（若存在）
    """
    if ocr_engine is not None:
        return ocr_engine.extract_text_from_pdf(pdf_path, max_pages)

    try:
        with pdfplumber.open(pdf_path) as pdf:
            metadata = pdf.metadata or {}
            texts = []
            front_text = ""

            for idx, page in enumerate(pdf.pages):
                if max_pages is not None and idx >= max_pages:
                    break

                page_text = page.extract_text() or ""
                page_text = page_text.strip()
                if page_text:
                    texts.append(page_text)
                    if not front_text:
                        front_text = page_text

            return "\n\n".join(texts), front_text, metadata
    except Exception as e:
        print(f"  错误: {e}")
        return "", "", {}


def _get_metadata_value(metadata, key_candidates):
    if not isinstance(metadata, dict):
        return ""
    for key in key_candidates:
        value = metadata.get(key)
        if value:
            return str(value).strip()
    # 一些 PDF 元数据键可能大小写不固定
    lowered = {str(k).lower(): v for k, v in metadata.items()}
    for key in key_candidates:
        value = lowered.get(str(key).lower())
        if value:
            return str(value).strip()
    return ""


def extract_title(front_text, filename, metadata):
    """智能提取标题"""
    # 从元数据提取（如果可靠）
    meta_title = _get_metadata_value(metadata, ["Title", "title"])
    if meta_title:
        lowered = meta_title.lower()
        if lowered not in {"", "untitled", "unknown"} and len(meta_title) >= 8:
            return " ".join(meta_title.split())

    # 从文件名提取
    if ' - ' in filename:
        parts = filename.split(' - ')
        if len(parts) >= 3:
            title_from_name = parts[-1].replace('.pdf', '').strip()
            if len(title_from_name) > 20:
                return title_from_name
    
    # 从文本提取
    raw_lines = front_text.split("\n")
    lines = [" ".join(l.strip().split()) for l in raw_lines[:120]]
    candidates = []
    keywords = [
        'quantum', 'qled', 'led', 'dot', 'efficiency', 'inp', 'passivation',
        'ligand', 'device', 'perovskite', 'solar',
    ]
    skip_tokens = ['©', 'http', 'www.', 'vol.', 'issn', 'available', 'doi', 'journal']
    affil_tokens = [
        "university", "institute", "department", "laboratory", "school",
        "college", "academy", "centre", "center",
    ]

    def is_candidate_title_line(line):
        if not line:
            return False
        lower = line.lower()
        if not (20 < len(line) < 260):
            return False
        if any(skip in lower for skip in skip_tokens):
            return False
        if re.search(r'\b(abstract|introduction|received|accepted|corresponding)\b', lower):
            return False
        # 过多数字/符号通常不是标题
        if sum(ch.isdigit() for ch in line) > 8:
            return False
        return True

    for i in range(min(80, len(lines))):
        line = lines[i]
        if not is_candidate_title_line(line):
            continue

        lower = line.lower()
        score = len(line)
        if any(kw in lower for kw in keywords):
            score += 40
        comma_count = line.count(",") + line.count("，") + line.count(";") + line.count("；")
        if comma_count:
            score -= comma_count * 15
        if '@' in line:
            score -= 40
        if any(tok in lower for tok in affil_tokens):
            score -= 60
        candidates.append((score, i, line))

        # 处理标题跨两行的情况（简单拼接）
        if i + 1 < len(lines):
            next_line = lines[i + 1]
            combined = f"{line} {next_line}".strip()
            next_lower = next_line.lower()
            next_commas = next_line.count(",") + next_line.count("，") + next_line.count(";") + next_line.count("；")
            if (
                is_candidate_title_line(next_line)
                and len(combined) < 260
                and next_commas <= 1
                and "@" not in next_line
                and not any(tok in next_lower for tok in affil_tokens)
            ):
                combined_lower = combined.lower()
                combined_score = len(combined) + 10
                if any(kw in combined_lower for kw in keywords):
                    combined_score += 40
                if any(tok in combined_lower for tok in affil_tokens):
                    combined_score -= 60
                candidates.append((combined_score, i, combined))

    if candidates:
        candidates.sort(key=lambda x: (-x[0], x[1]))
        return candidates[0][2]
    
    return "未提取到标题"

def extract_journal_info(front_text, impact_factors):
    """从PDF首页/前置信息中提取期刊信息（优先匹配影响因子库中的期刊名）"""
    text = (front_text or "")[:8000]

    # 先匹配常见缩写
    alias_patterns = [
        (r"\bJ\.?\s*Am\.?\s*Chem\.?\s*Soc\.?\b", "Journal of the American Chemical Society"),
        (r"\bPhys\.?\s*Chem\.?\s*Chem\.?\s*Phys\.?\b", "Physical Chemistry Chemical Physics"),
        (r"\bJ\.?\s*Mater\.?\s*Chem\.?\s*C\b", "Journal of Materials Chemistry C"),
        (r"\bAdv\.?\s*Mater\.?\b", "Advanced Materials"),
        (r"\bAdv\.?\s*Funct\.?\s*Mater\.?\b", "Advanced Functional Materials"),
        (r"\bNano\s*Lett\.?\b", "Nano Letters"),
        (r"\bNat\.?\s*Commun\.?\b", "Nature Communications"),
    ]
    for pattern, canonical in alias_patterns:
        if re.search(pattern, text, re.IGNORECASE):
            return canonical

    # 再从影响因子库中做最长优先匹配
    journals = sorted(impact_factors.keys(), key=lambda s: len(str(s)), reverse=True)
    best = None
    for journal in journals:
        if not journal:
            continue
        escaped = re.escape(str(journal))
        pattern = re.compile(rf"(?<![A-Za-z0-9]){escaped}(?![A-Za-z0-9])", re.IGNORECASE)
        match = pattern.search(text)
        if not match:
            continue

        candidate = (match.start(), -len(str(journal)), -float(impact_factors.get(journal, 0.0)), str(journal))
        if best is None or candidate < best:
            best = candidate

    if best:
        return best[3]

    return "未知期刊"


def extract_authors(front_text, metadata):
    """从PDF文本/元数据中提取作者信息"""
    meta_author = _get_metadata_value(metadata, ["Author", "author"])
    if meta_author and len(meta_author) >= 3:
        parts = [p.strip() for p in re.split(r"[;,\n]+", meta_author) if p.strip()]
        if parts:
            if len(parts) > 3:
                return ", ".join(parts[:3]) + " et al."
            return ", ".join(parts[:3])

    lines = (front_text or "")[:8000].split("\n")
    authors = []

    name_regex = re.compile(r"\b[A-Z][a-zA-Z\-']+\s+[A-Z][a-zA-Z\-']+\b")
    skip_words = [
        "journal", "doi", "http", "www", "©", "vol.", "issn", "abstract",
        "introduction", "received", "accepted", "available", "corresponding",
    ]

    best_single = ""
    for line in lines[:100]:
        raw = line.strip()
        if not raw or len(raw) < 5:
            continue
        lower = raw.lower()
        if any(sw in lower for sw in skip_words):
            continue

        cleaned = re.sub(r"[\*\†\‡\§\¶]", "", raw)
        cleaned = re.sub(r"\d+", "", cleaned)
        cleaned = " ".join(cleaned.split())

        matches = name_regex.findall(cleaned)
        # 标题行可能误命中 1 个“像姓名的词组”，作者行通常 >=2
        if len(matches) >= 2:
            authors = matches
            break
        if not best_single and len(matches) == 1:
            # 兜底：仅 1 个作者的情况，或分行导致只提取到 1 个姓名
            if 10 < len(cleaned) < 120:
                best_single = matches[0]
    
    if authors:
        if len(authors) > 3:
            return ', '.join(authors[:3]) + ' et al.'
        return ', '.join(authors)

    if best_single:
        return best_single
    
    return "未提取"

def get_impact_factor(journal_name, impact_factors):
    """获取期刊影响因子"""
    return impact_factors.get(journal_name, 0.0)

def extract_eqe_comprehensive(text):
    """全面提取EQE"""
    values = []
    
    patterns = [
        r'EQE[^0-9<≥>]*?([0-9]+\.?[0-9]*)\s*%',
        r'external quantum efficiency[^0-9<≥>]*?([0-9]+\.?[0-9]*)\s*%',
        r'外(?:部)?量子效率[^0-9<≥>]*?([0-9]+\.?[0-9]*)\s*%',
        r'η\s*[_ ]?ext[^0-9<≥>]*?([0-9]+\.?[0-9]*)\s*%',
        r'EQE[^0-9]*?[≥>]\s*([0-9]+\.?[0-9]*)\s*%',
        r'maximum EQE[^0-9]*?([0-9]+\.?[0-9]*)\s*%',
        r'peak EQE[^0-9]*?([0-9]+\.?[0-9]*)\s*%',
    ]
    
    for pattern in patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        for m in matches:
            try:
                v = float(m)
                if 0.1 < v < 80:
                    values.append(v)
            except:
                pass
    
    if values:
        return f"{max(values):.2f}%"
    return ""

def extract_cie_comprehensive(text):
    """全面提取CIE坐标"""
    patterns = [
        r'CIE[^0-9]*?\(([0-9]\.[0-9]+)\s*[,，]\s*([0-9]\.[0-9]+)\)',
        r'CIE[^0-9]*?([0-9]\.[0-9]+)\s*[,，]\s*([0-9]\.[0-9]+)',
        r'chromaticity[^0-9]*?\(([0-9]\.[0-9]+)\s*[,，]\s*([0-9]\.[0-9]+)\)',
    ]
    
    for pattern in patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        for m in matches:
            try:
                x, y = float(m[0]), float(m[1])
                if 0 < x < 1 and 0 < y < 1:
                    return f"({x:.4f}, {y:.4f})"
            except:
                pass
    
    return ""

def extract_lifetime_comprehensive(text):
    """全面提取寿命"""
    values = []
    
    patterns = [
        r'T[⑤5]0[^0-9]*?([0-9]+\.?[0-9]*)\s*(h|hr|hrs|hour|hours|小时)',
        r'\bT\s*50\b[^0-9]*?([0-9]+\.?[0-9]*)\s*(h|hr|hrs|hour|hours|小时)',
        r'LT[⑤5]0[^0-9]*?([0-9]+\.?[0-9]*)\s*(h|hr|hrs|hour|hours|小时)',
        r'\bLT\s*50\b[^0-9]*?([0-9]+\.?[0-9]*)\s*(h|hr|hrs|hour|hours|小时)',
        r'lifetime[^0-9]*?([0-9]+\.?[0-9]*)\s*(h|hour|hours|小时)',
        r'operational[^0-9]*?([0-9]+\.?[0-9]*)\s*(h|hour|hours)',
    ]
    
    for pattern in patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        for m in matches:
            try:
                v = float(m[0])
                if 1 < v < 50000:
                    values.append(v)
            except:
                pass
    
    if values:
        max_v = max(values)
        if max_v >= 1000:
            return f"{max_v:.0f} h ({max_v/1000:.2f}k h)"
        return f"{max_v:.1f} h"
    return ""

def extract_device_structure(text):
    """提取器件结构"""
    materials = []
    
    layer_keywords = {
        'ITO': r'\bITO\b',
        'PEDOT:PSS': r'PEDOT:PSS|PEDOT\s*:\s*PSS',
        'TFB': r'\bTFB\b',
        'PVK': r'\bPVK\b',
        'Poly-TPD': r'Poly-TPD|Poly\s*-\s*TPD',
        'QD层': r'quantum dot|QD|InP\.?QD|CdSe',
        'ZnO': r'\bZnO\b',
        'ZnMgO': r'ZnMgO|Zn\d*Mg\d*O',
        'TPBi': r'\bTPBi\b',
        'LiF': r'\bLiF\b',
        'Al': r'\bAl\b(?!.*\bGa\b)',
        'Ag': r'\bAg\b',
        'MoO3': r'MoO\d|MoO3',
        'NiO': r'\bNiO\b',
    }
    
    for mat, pattern in layer_keywords.items():
        if re.search(pattern, text, re.IGNORECASE):
            materials.append(mat)
    
    order = ['ITO', 'PEDOT:PSS', 'Poly-TPD', 'TFB', 'PVK', 'QD层', 'ZnO', 'ZnMgO', 'TPBi', 'LiF', 'Al', 'Ag', 'MoO3', 'NiO']
    
    ordered = []
    for mat in order:
        if mat in materials:
            ordered.append(mat)
    
    return ' / '.join(ordered) if ordered else ""

def extract_optimization(text):
    """提取优化信息"""
    text_lower = text.lower()
    
    levels = []
    level_keywords = {
        '材料合成': ['synthesis', 'material design', 'precursor'],
        '核壳结构': ['core-shell', 'core/shell', 'shell growth'],
        '表面处理': ['surface treatment', 'surface modification', 'passivation'],
        '配体工程': ['ligand engineering', 'ligand exchange', 'ligand modification'],
        '器件结构': ['device architecture', 'device structure', 'interface engineering'],
        '工艺优化': ['annealing', 'thermal treatment', 'solution processing'],
    }
    
    for level, keywords in level_keywords.items():
        if any(kw in text_lower for kw in keywords):
            levels.append(level)
    
    strategies = []
    strategy_keywords = {
        '表面钝化': ['passivation', 'defect passivation'],
        '配体交换': ['ligand exchange', 'ligand replacement'],
        '核壳工程': ['core-shell', 'shell growth'],
        '界面工程': ['interface engineering', 'interface modification'],
        '退火处理': ['annealing', 'thermal treatment'],
        '溶剂工程': ['solvent engineering', 'solvent treatment'],
        '元素掺杂': ['doping', 'doped', 'dopant'],
        '刻蚀策略': ['etching', 'core etching'],
    }
    
    for strategy, keywords in strategy_keywords.items():
        if any(kw in text_lower for kw in keywords):
            strategies.append(strategy)
    
    return '、'.join(levels), '、'.join(strategies)

def extract_supplementary(text):
    """提取补充信息"""
    info = []
    
    # 波长
    wl_patterns = [
        r'emission wavelength[^0-9]*?([0-9]{3})\s*nm',
        r'peak wavelength[^0-9]*?([0-9]{3})\s*nm',
        r'([0-9]{3})\s*nm[^0-9]*?(?:emission|PL|peak)',
    ]
    
    wavelengths = []
    for pattern in wl_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        for m in matches:
            try:
                wl = int(m)
                if 400 <= wl <= 800:
                    wavelengths.append(wl)
            except:
                pass
    
    if wavelengths:
        min_wl, max_wl = min(wavelengths), max(wavelengths)
        avg_wl = sum(wavelengths) / len(wavelengths)
        
        if 450 <= avg_wl < 495:
            color = "蓝光"
        elif 495 <= avg_wl < 570:
            color = "绿光"
        elif 570 <= avg_wl < 620:
            color = "黄光"
        elif 620 <= avg_wl <= 750:
            color = "红光"
        else:
            color = ""
        
        if min_wl == max_wl:
            info.append(f"发射波长: {min_wl} nm")
        else:
            info.append(f"发射波长: {min_wl}-{max_wl} nm")
        
        if color:
            info.append(f"颜色: {color}")
    
    # 亮度
    lum_patterns = [
        r'luminance[^0-9]*?([0-9,]+)\s*cd/m',
        r'brightness[^0-9]*?([0-9,]+)\s*cd/m',
        r'luminance[^0-9]*?([0-9,]+)\s*cd\s*m',
        r'brightness[^0-9]*?([0-9,]+)\s*cd\s*m',
    ]
    
    for pattern in lum_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        if matches:
            try:
                lum = int(matches[0].replace(',', ''))
                if lum > 100:
                    info.append(f"亮度: {lum:,} cd/m²")
                    break
            except:
                pass
    
    return '; '.join(info)


def infer_paper_type(title, full_text, front_text):
    """粗略判断文章类型，用于解释“指标缺失”的可能原因。

    说明：这是启发式判断，目的是把“非器件性能类论文”与“器件论文”区分开，
    以便解释 EQE/CIE/寿命 等指标缺失时，是否更可能属于“文章类型原因”。
    """
    sample = "\n".join(
        [
            str(title or ""),
            str(front_text or ""),
            str(full_text or "")[:8000],
        ]
    )
    sample_lower = sample.lower()
    evidence = []

    review_patterns = [
        r"\breview\b",
        r"\bperspective\b",
        r"\boverview\b",
        r"\bprogress\b",
        r"\badvances?\b",
        r"\broadmap\b",
        r"\bsurvey\b",
        r"综述",
        r"进展",
        r"展望",
        r"回顾",
    ]
    if any(re.search(pat, sample_lower, re.IGNORECASE) for pat in review_patterns):
        evidence.append("命中综述/展望关键词")
        return "综述/展望", evidence

    theory_patterns = [
        r"\btheoretical\b",
        r"\btheory\b",
        r"\bsimulation\b",
        r"\bsimulated\b",
        r"\bcomputational\b",
        r"\bdft\b",
        r"first[- ]principles",
        r"密度泛函",
        r"第一性原理",
        r"模拟",
        r"仿真",
    ]
    if any(re.search(pat, sample_lower, re.IGNORECASE) for pat in theory_patterns):
        evidence.append("命中理论/计算关键词")
        return "理论/计算", evidence

    device_keywords = [
        "device",
        "qled",
        "oled",
        "led",
        "electroluminescence",
        "external quantum efficiency",
        " eqe",
        "luminance",
        "cd/m",
        "t50",
        "lt50",
        "lifetime",
        "turn-on voltage",
        "current efficiency",
        "power efficiency",
        "solar cell",
        "photovoltaic",
        "pce",
    ]
    hits = []
    for kw in device_keywords:
        if kw.strip() and kw in sample_lower:
            hits.append(kw.strip())
    unique_hits = sorted(set(hits))
    if len(unique_hits) >= 4:
        evidence.append(f"器件关键词命中: {', '.join(unique_hits[:8])}")
        return "器件/性能", evidence

    if unique_hits:
        evidence.append(f"弱器件信号: {', '.join(unique_hits[:6])}")
    return "材料/机理/其他", evidence


def _field_is_missing(field_name, value):
    if value is None:
        return True
    if isinstance(value, str):
        v = value.strip()
        if not v:
            return True
        if field_name == "期刊名称" and v == "未知期刊":
            return True
        if field_name == "作者" and v in {"未提取", "未提取到作者"}:
            return True
        if field_name == "论文标题" and v in {"未提取到标题", "未提取"}:
            return True
    return False


def _find_keyword_hits(text_lower, keyword_patterns):
    hits = []
    for label, pattern in keyword_patterns:
        if re.search(pattern, text_lower, re.IGNORECASE):
            hits.append(label)
    return hits


def analyze_missing_items(result, paper_type, full_text, front_text):
    """分析缺失项，并给出“文章类型原因/可能未给出/提及但未提取”等解释。"""
    if not result:
        return []

    text_lower = (full_text or "").lower()
    if front_text:
        text_lower = f"{front_text}\n{text_lower}".lower()

    keyword_fields = {
        "EQE（外量子效率）": [
            ("EQE", r"\beqe\b"),
            ("external quantum efficiency", r"external\s+quantum\s+efficiency"),
            ("η_ext", r"η\s*[_ ]?ext"),
            ("外(部)量子效率", r"外(?:部)?量子效率"),
        ],
        "色度坐标": [
            ("CIE", r"\bcie\b"),
            ("chromaticity", r"chromaticity"),
            ("(x,y)", r"\(\s*0\.\d+\s*[,，]\s*0\.\d+\s*\)"),
            ("色度", r"色度"),
        ],
        "寿命": [
            ("T50", r"\bt\s*50\b|\bt50\b"),
            ("LT50", r"\blt\s*50\b|\blt50\b"),
            ("lifetime", r"\blifetime\b"),
            ("operational", r"\boperational\b"),
            ("寿命", r"寿命|半衰"),
        ],
        "器件结构": [
            ("ITO", r"\bito\b"),
            ("PEDOT:PSS", r"pedot\s*:\s*pss|pedot:p\w*s"),
            ("TFB", r"\btfb\b"),
            ("PVK", r"\bpvk\b"),
            ("Poly-TPD", r"poly\s*-\s*tpd|poly-tpd"),
            ("ZnO", r"\bzno\b"),
            ("ZnMgO", r"zn\s*\d*\s*mg\s*\d*\s*o|znmgo"),
            ("TPBi", r"\btpbi\b"),
            ("LiF", r"\blif\b"),
            ("Al", r"\bal\b"),
            ("Ag", r"\bag\b"),
        ],
        "优化层级": [
            ("surface", r"surface\s+(treatment|modification)|passivation"),
            ("ligand", r"ligand\s+(engineering|exchange|replacement)"),
            ("core-shell", r"core[-/ ]shell|shell\s+growth"),
            ("interface", r"interface\s+engineering|interface\s+modification"),
            ("anneal", r"anneal|annealing|thermal\s+treatment"),
            ("solvent", r"solvent\s+engineering|solvent\s+treatment"),
        ],
        "优化策略": [
            ("passivation", r"passivation|defect\s+passivation"),
            ("ligand exchange", r"ligand\s+exchange|ligand\s+replacement"),
            ("core-shell", r"core[-/ ]shell|shell\s+growth"),
            ("interface", r"interface\s+engineering|interface\s+modification"),
            ("annealing", r"anneal|annealing|thermal\s+treatment"),
            ("doping", r"\bdoping\b|\bdoped\b|\bdopant\b"),
            ("etching", r"etching|core\s+etching"),
        ],
    }

    missing_items = []

    for field in REPORT_HEADERS:
        if field in {"File", "URL"}:
            continue
        value = result.get(field, "")
        if not _field_is_missing(field, value):
            continue

        reason = "未提取/未识别"
        evidence = ""

        if field in keyword_fields:
            if paper_type != "器件/性能" and field in {"EQE（外量子效率）", "色度坐标", "寿命", "器件结构"}:
                reason = "文章类型原因（非器件性能类论文，指标可能不适用）"
            else:
                hits = _find_keyword_hits(text_lower, keyword_fields[field])
                if hits:
                    if field in {"EQE（外量子效率）", "色度坐标", "寿命"}:
                        reason = "文本提及该指标，但未匹配到数值（可能格式不同）"
                    else:
                        reason = "文本出现相关关键词，但未被规则识别（可能表达方式不同）"
                    evidence = "关键词命中: " + ", ".join(hits[:4])
                else:
                    if field in {"EQE（外量子效率）", "色度坐标", "寿命"} and paper_type == "器件/性能":
                        reason = "未发现相关指标关键词（更可能文章确实没有/未给出）"
                    elif field == "器件结构" and paper_type == "器件/性能":
                        reason = "未发现常见器件层材料关键词（更可能文章未明确给出器件结构）"
                    elif field in {"优化层级", "优化策略"}:
                        reason = "未发现常见优化关键词（更可能文章未强调优化策略/层级）"

        missing_items.append(
            {
                "field": field,
                "reason": reason,
                "evidence": evidence,
            }
        )

    return missing_items


def analyze_pdf_record(pdf_path, impact_factors, max_pages: Optional[int] = None, ocr_engine=None):
    """分析单个PDF文件，返回:
    - result: 用于 Excel/JSON/CSV 的结果 dict（成功时）
    - record: 用于 MD/DOCX 缺失项报告的诊断信息（始终返回）
    """
    print(f"\n分析: {pdf_path.name}")

    text, front_text, metadata = extract_pdf_text(pdf_path, max_pages=max_pages, ocr_engine=ocr_engine)
    if not text:
        print("  警告: 未提取到文本（可能是扫描版PDF或无文本层）。")
        record = {
            "File": pdf_path.name,
            "URL": pdf_path.resolve().as_uri(),
            "status": "no_text",
            "paper_type": "无法判断",
            "paper_type_evidence": ["未提取到文本"],
            "missing_items": [
                {
                    "field": "全文文本",
                    "reason": "提取失败/无文本层（扫描版PDF需先OCR）",
                    "evidence": "",
                }
            ],
            "result": None,
        }
        return None, record

    title = extract_title(front_text or text, pdf_path.name, metadata)
    journal = extract_journal_info(front_text or text, impact_factors)
    authors = extract_authors(front_text or text, metadata)
    if_value = get_impact_factor(journal, impact_factors)
    device_structure = extract_device_structure(text)
    opt_level, opt_strategy = extract_optimization(text)
    eqe = extract_eqe_comprehensive(text)
    cie = extract_cie_comprehensive(text)
    lifetime = extract_lifetime_comprehensive(text)
    supp_info = extract_supplementary(text)

    print(f"  期刊: {journal} (IF: {if_value})")
    print(f"  EQE: {eqe}")
    print(f"  寿命: {lifetime}")

    result = {
        "File": pdf_path.name,
        "URL": pdf_path.resolve().as_uri(),
        "期刊名称": journal,
        "影响因子": if_value,
        "作者": authors,
        "论文标题": title,
        "器件结构": device_structure,
        "优化层级": opt_level,
        "优化策略": opt_strategy,
        "EQE（外量子效率）": eqe,
        "色度坐标": cie,
        "寿命": lifetime,
        "补充信息": supp_info,
    }

    paper_type, paper_type_evidence = infer_paper_type(title, text, front_text)
    missing_items = analyze_missing_items(result, paper_type, text, front_text)

    record = {
        "File": pdf_path.name,
        "URL": pdf_path.resolve().as_uri(),
        "status": "ok",
        "paper_type": paper_type,
        "paper_type_evidence": paper_type_evidence,
        "missing_items": missing_items,
        "result": result,
    }
    return result, record


def sort_results(results):
    """按影响因子降序排序（None 安全）"""
    return sorted(results, key=lambda x: x.get("影响因子", 0) if x else 0, reverse=True)


def create_excel_report(results, output_path):
    """创建Excel报告"""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "论文分析结果"
    
    # 表头
    headers = list(REPORT_HEADERS)
    highlight_headers = {"EQE（外量子效率）", "色度坐标", "寿命"}
    highlight_cols = {headers.index(h) + 1 for h in highlight_headers if h in headers}
    
    # 样式
    header_fill = PatternFill(start_color="2E75B6", end_color="2E75B6", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=11, name='Arial')
    cell_font = Font(size=10, name='Arial')
    thin_border = Border(
        left=Side(style='thin', color='BFBFBF'),
        right=Side(style='thin', color='BFBFBF'),
        top=Side(style='thin', color='BFBFBF'),
        bottom=Side(style='thin', color='BFBFBF')
    )
    
    # 写入表头
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        cell.border = thin_border
    
    # 按影响因子排序
    results_sorted = sort_results(results)
    
    # 写入数据
    for row_idx, result in enumerate(results_sorted, 2):
        if result:
            for col_idx, header in enumerate(headers, 1):
                value = result.get(header, '')
                cell = ws.cell(row=row_idx, column=col_idx, value=value)
                cell.font = cell_font
                cell.alignment = Alignment(vertical='top', wrap_text=True)
                cell.border = thin_border

                if header == "URL" and isinstance(value, str) and value.startswith("file://"):
                    cell.hyperlink = value
                    cell.font = Font(size=10, name="Arial", color="0563C1", underline="single")
                
                # 高亮有数据的单元格
                if col_idx in highlight_cols and value:
                    cell.fill = PatternFill(start_color="E2EFDA", end_color="E2EFDA", fill_type="solid")
    
    # 设置列宽
    column_widths = [30, 45, 25, 12, 30, 55, 45, 30, 35, 18, 22, 22, 40]
    for i, width in enumerate(column_widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = width
    
    # 设置行高
    ws.row_dimensions[1].height = 35
    for row in range(2, len(results) + 2):
        ws.row_dimensions[row].height = 80
    
    # 冻结首行
    ws.freeze_panes = 'A2'
    
    # 保存
    wb.save(output_path)
    print(f"\n✓ Excel报告已保存: {output_path}")


def export_json_report(results, output_path):
    """导出 JSON 报告（UTF-8）"""
    results_sorted = sort_results(results)
    with output_path.open("w", encoding="utf-8") as f:
        json.dump(results_sorted, f, ensure_ascii=False, indent=2)
    print(f"✓ JSON结果已保存: {output_path}")


def export_csv_report(results, output_path):
    """导出 CSV 报告（UTF-8-SIG，方便 Excel/WPS 直接打开）"""
    results_sorted = sort_results(results)
    headers = list(REPORT_HEADERS)
    with output_path.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=headers, extrasaction="ignore")
        writer.writeheader()
        for row in results_sorted:
            writer.writerow({h: row.get(h, "") for h in headers})
    print(f"✓ CSV结果已保存: {output_path}")


def _classify_missing_reason(reason):
    if not reason:
        return "其他"
    if reason.startswith("文章类型原因"):
        return "文章类型原因"
    if reason.startswith("未发现相关指标关键词"):
        return "文章未给出/不涉及"
    if reason.startswith("文本提及该指标"):
        return "提及但未提取"
    if reason.startswith("提取失败/无文本层"):
        return "提取失败/无文本层"
    if reason.startswith("未提取/未识别"):
        return "未提取/未识别"
    return "其他"


def export_markdown_missing_report(records, output_path, run_info):
    """导出 Markdown 缺失项报告"""
    def md_escape(value):
        s = "" if value is None else str(value)
        s = s.replace("\r", " ").replace("\n", " ").strip()
        s = s.replace("|", "\\|")
        return s

    generated_at = run_info.get("generated_at", "")
    max_pages = run_info.get("max_pages")
    pdf_count = len(records)
    ok_count = sum(1 for r in records if r.get("status") == "ok")
    fail_count = pdf_count - ok_count

    field_reason_counts = {}
    for rec in records:
        for item in rec.get("missing_items", []) or []:
            field = item.get("field", "")
            reason = item.get("reason", "")
            bucket = _classify_missing_reason(reason)
            field_reason_counts.setdefault(field, {})
            field_reason_counts[field][bucket] = field_reason_counts[field].get(bucket, 0) + 1

    reason_cols = [
        "文章类型原因",
        "文章未给出/不涉及",
        "提及但未提取",
        "提取失败/无文本层",
        "未提取/未识别",
        "其他",
    ]

    lines = []
    lines.append("# 论文缺失项报告")
    lines.append("")
    if generated_at:
        lines.append(f"- 生成时间: {generated_at}")
    lines.append(f"- 总PDF数: {pdf_count}")
    lines.append(f"- 成功解析: {ok_count}")
    lines.append(f"- 无文本层/失败: {fail_count}")
    if max_pages is not None:
        lines.append(f"- max_pages: {max_pages}（仅基于已读取页数判断缺失）")
    lines.append("- 说明: “缺失原因”基于关键词与文章类型的启发式判断，仅供参考。")
    lines.append("")

    if field_reason_counts:
        lines.append("## 缺失汇总")
        lines.append("")
        lines.append("| 字段 | 缺失数 | " + " | ".join(reason_cols) + " |")
        lines.append("|---|---:|" + "|".join(["---:"] * len(reason_cols)) + "|")
        for field in sorted(field_reason_counts.keys(), key=lambda f: (-sum(field_reason_counts[f].values()), f)):
            total = sum(field_reason_counts[field].values())
            row = [field, str(total)]
            for col in reason_cols:
                row.append(str(field_reason_counts[field].get(col, 0)))
            lines.append("| " + " | ".join(row) + " |")
        lines.append("")

    lines.append("## 逐篇缺失项")
    lines.append("")
    for rec in records:
        file_name = rec.get("File", "")
        url = rec.get("URL", "")
        status = rec.get("status", "")
        paper_type = rec.get("paper_type", "")
        type_evidence = rec.get("paper_type_evidence") or []
        missing_items = rec.get("missing_items") or []
        res = rec.get("result") or {}

        lines.append(f"### {md_escape(file_name)}")
        if url:
            lines.append(f"- URL: {md_escape(url)}")
        lines.append(f"- 状态: {md_escape(status)}")
        if status == "ok":
            journal = res.get("期刊名称", "")
            if_value = res.get("影响因子", "")
            title = res.get("论文标题", "")
            authors = res.get("作者", "")
            lines.append(f"- 期刊: {md_escape(journal)} (IF: {md_escape(if_value)})")
            if title:
                lines.append(f"- 标题: {md_escape(title)}")
            if authors:
                lines.append(f"- 作者: {md_escape(authors)}")
            if paper_type:
                evidence = "; ".join([str(e).strip() for e in type_evidence if str(e).strip()])
                if evidence:
                    lines.append(f"- 类型判断: {md_escape(paper_type)}（{md_escape(evidence)}）")
                else:
                    lines.append(f"- 类型判断: {md_escape(paper_type)}")
        else:
            if paper_type:
                lines.append(f"- 类型判断: {md_escape(paper_type)}")

        if not missing_items:
            lines.append("- 缺失项: 无")
            lines.append("")
            continue

        lines.append("")
        lines.append("| 缺失字段 | 原因 | 线索 |")
        lines.append("|---|---|---|")
        for item in missing_items:
            field = md_escape(item.get("field", ""))
            reason = md_escape(item.get("reason", ""))
            evidence = md_escape(item.get("evidence", ""))
            lines.append(f"| {field} | {reason} | {evidence} |")
        lines.append("")

    with output_path.open("w", encoding="utf-8") as f:
        f.write("\n".join(lines).rstrip() + "\n")
    print(f"✓ Markdown缺失项报告已保存: {output_path}")


def export_docx_missing_report(records, output_path, run_info):
    """导出 Word(.docx) 缺失项报告"""
    try:
        from docx import Document
        from docx.oxml.ns import qn
    except ImportError as exc:
        print(f"[警告] 未安装 python-docx，无法导出 DOCX: {exc}")
        return

    generated_at = run_info.get("generated_at", "")
    max_pages = run_info.get("max_pages")
    pdf_count = len(records)
    ok_count = sum(1 for r in records if r.get("status") == "ok")
    fail_count = pdf_count - ok_count

    doc = Document()
    # 设置默认字体（兼容中英文）
    try:
        normal_style = doc.styles["Normal"]
        normal_style.font.name = "Arial"
        normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    except Exception:
        pass

    doc.add_heading("论文缺失项报告", level=0)
    if generated_at:
        doc.add_paragraph(f"生成时间: {generated_at}")
    doc.add_paragraph(f"总PDF数: {pdf_count}  成功解析: {ok_count}  无文本层/失败: {fail_count}")
    if max_pages is not None:
        doc.add_paragraph(f"max_pages: {max_pages}（仅基于已读取页数判断缺失）")
    doc.add_paragraph("说明: “缺失原因”基于关键词与文章类型的启发式判断，仅供参考。")

    # 汇总表
    field_reason_counts = {}
    for rec in records:
        for item in rec.get("missing_items", []) or []:
            field = item.get("field", "")
            reason = item.get("reason", "")
            bucket = _classify_missing_reason(reason)
            field_reason_counts.setdefault(field, {})
            field_reason_counts[field][bucket] = field_reason_counts[field].get(bucket, 0) + 1

    reason_cols = [
        "文章类型原因",
        "文章未给出/不涉及",
        "提及但未提取",
        "提取失败/无文本层",
        "未提取/未识别",
        "其他",
    ]

    if field_reason_counts:
        doc.add_heading("缺失汇总", level=1)
        table = doc.add_table(rows=1, cols=2 + len(reason_cols))
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "字段"
        hdr_cells[1].text = "缺失数"
        for idx, col in enumerate(reason_cols, start=2):
            hdr_cells[idx].text = col

        for field in sorted(field_reason_counts.keys(), key=lambda f: (-sum(field_reason_counts[f].values()), f)):
            total = sum(field_reason_counts[field].values())
            row_cells = table.add_row().cells
            row_cells[0].text = str(field)
            row_cells[1].text = str(total)
            for idx, col in enumerate(reason_cols, start=2):
                row_cells[idx].text = str(field_reason_counts[field].get(col, 0))

    doc.add_heading("逐篇缺失项", level=1)
    for rec in records:
        file_name = rec.get("File", "")
        url = rec.get("URL", "")
        status = rec.get("status", "")
        paper_type = rec.get("paper_type", "")
        type_evidence = rec.get("paper_type_evidence") or []
        missing_items = rec.get("missing_items") or []
        res = rec.get("result") or {}

        doc.add_heading(str(file_name), level=2)
        if url:
            doc.add_paragraph(f"URL: {url}")
        doc.add_paragraph(f"状态: {status}")
        if status == "ok":
            journal = res.get("期刊名称", "")
            if_value = res.get("影响因子", "")
            title = res.get("论文标题", "")
            authors = res.get("作者", "")
            doc.add_paragraph(f"期刊: {journal} (IF: {if_value})")
            if title:
                doc.add_paragraph(f"标题: {title}")
            if authors:
                doc.add_paragraph(f"作者: {authors}")
            if paper_type:
                ev = "; ".join([str(e).strip() for e in type_evidence if str(e).strip()])
                if ev:
                    doc.add_paragraph(f"类型判断: {paper_type}（{ev}）")
                else:
                    doc.add_paragraph(f"类型判断: {paper_type}")
        else:
            if paper_type:
                doc.add_paragraph(f"类型判断: {paper_type}")

        if not missing_items:
            doc.add_paragraph("缺失项: 无")
            continue

        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "缺失字段"
        hdr[1].text = "原因"
        hdr[2].text = "线索"
        for item in missing_items:
            row = table.add_row().cells
            row[0].text = str(item.get("field", ""))
            row[1].text = str(item.get("reason", ""))
            row[2].text = str(item.get("evidence", ""))

    doc.save(output_path)
    print(f"✓ Word缺失项报告已保存: {output_path}")


def main():
    """主函数"""
    args = parse_args()
    pdf_dir = args.pdf_dir.expanduser().resolve()
    output_dir = args.output_dir.expanduser().resolve()
    journal_config_path = args.journal_config.expanduser().resolve()
    use_existing_json = USE_EXISTING_JSON and not args.no_json
    impact_factors = load_journal_impact_factors(
        journal_config_path,
        use_existing_json=use_existing_json,
    )

    ocr_engine = None
    use_ocr = args.use_ocr and not args.no_ocr
    if use_ocr:
        print("\n[OCR] 启用 PaddleOCR 文本提取模式")
        ocr_engine = create_ocr_engine(
            lang=args.ocr_lang,
            enable_hpi=args.enable_hpi,
        )
    else:
        print("\n[INFO] 使用 pdfplumber 提取文本（可通过 --use-ocr 启用OCR）")

    print("="*70)
    print(" PDF论文分析工具 v1.4 (OCR增强版)")
    print("="*70)
    
    # 创建输出目录
    output_dir.mkdir(parents=True, exist_ok=True)
    
    # 获取PDF文件
    pdf_files = iter_pdf_files(pdf_dir, recursive=args.recursive)
    max_pages = args.max_pages if args.max_pages and args.max_pages > 0 else None
    print(f"\n找到 {len(pdf_files)} 个PDF文件")
    
    if not pdf_files:
        print("\n错误: 未找到PDF文件！")
        if pdf_dir.is_file():
            print(f"请确认输入路径是 PDF 文件: {pdf_dir}")
        else:
            print(f"请将PDF文件放入: {pdf_dir}")
        return
    
    # 分析每个PDF
    results = []
    records = []
    for i, pdf_file in enumerate(pdf_files, 1):
        print(f"\n[{i}/{len(pdf_files)}]", end='')
        result, record = analyze_pdf_record(pdf_file, impact_factors, max_pages=max_pages, ocr_engine=ocr_engine)
        records.append(record)
        if result:
            results.append(result)
    
    # 生成报告
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    report_base = f"论文分析报告_{timestamp}"
    missing_base = f"论文缺失项报告_{timestamp}"

    if results:
        output_file = output_dir / f"{report_base}.xlsx"
        create_excel_report(results, output_file)
    else:
        print("\n错误: 所有PDF均未成功提取到文本，未生成Excel报告。")

    if args.export_json:
        if results:
            export_json_report(results, output_dir / f"{report_base}.json")
        else:
            print("[提示] 无可用结果，跳过 JSON 导出。")
    if args.export_csv:
        if results:
            export_csv_report(results, output_dir / f"{report_base}.csv")
        else:
            print("[提示] 无可用结果，跳过 CSV 导出。")
    run_info = {
        "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "max_pages": max_pages,
    }
    if args.export_md:
        export_markdown_missing_report(records, output_dir / f"{missing_base}.md", run_info)
    if args.export_docx:
        export_docx_missing_report(records, output_dir / f"{missing_base}.docx", run_info)
    
    # 统计
    eqe_count = sum(1 for r in results if r["EQE（外量子效率）"])
    lifetime_count = sum(1 for r in results if r["寿命"])
    cie_count = sum(1 for r in results if r["色度坐标"])
    
    print(f"\n{'='*70}")
    print(" 分析完成!")
    print(f"{'='*70}")
    print(f" 总文件数: {len(results)}/{len(pdf_files)}")
    if results:
        print(f" EQE数据: {eqe_count}/{len(results)}")
        print(f" 寿命数据: {lifetime_count}/{len(results)}")
        print(f" CIE坐标: {cie_count}/{len(results)}")
    else:
        print(" EQE数据: 0/0")
        print(" 寿命数据: 0/0")
        print(" CIE坐标: 0/0")
    print(f"{'='*70}")

if __name__ == "__main__":
    main()
